import streamlit as st
import fitz
from google import genai
import io
import re
import os
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

# [설정 - API]
MY_API_KEY = st.secrets["GEMINI_API_KEY"].strip() 
client = genai.Client(api_key=MY_API_KEY)

# 고정 템플릿 파일명 설정 (같은 폴더 내에 있어야 함)
DEFAULT_WORD_TEMPLATE = "default_vf_template.docx"

# 1. 폰트 설정
def set_font(run, font_name, size, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.bold = bold

# 2. PDF 텍스트 추출
def extract_text_from_pdf(uploaded_file):
    if uploaded_file is None: return ""
    uploaded_file.seek(0)
    doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
    return "".join([page.get_text() for page in doc])[:15000]

# 3. 특정 단락 뒤에 스타일이 적용된 콘텐츠를 삽입하는 함수 (템플릿용)
def add_styled_content_at(target_p, text):
    lines = str(text).split('\n')
    current_p = target_p
    for line in lines:
        line_stripped = line.strip()
        if not line_stripped: continue
        
        new_p_xml = OxmlElement('w:p')
        current_p._p.addnext(new_p_xml)
        new_p = Paragraph(new_p_xml, current_p._parent)
        
        if line_stripped.startswith('## '):
            run = new_p.add_run(line_stripped.replace('## ', ''))
            set_font(run, "KoPub돋움체_Pro Medium", 12, bold=True)
        else:
            new_p.paragraph_format.line_spacing = 1.6
            new_p.paragraph_format.space_after = Pt(12)
            
            parts = re.split(r'(\*\*.*?\*\*)', line_stripped)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = new_p.add_run(part.replace('**', ''))
                    set_font(run, "KoPub돋움체_Pro Medium", 11, bold=True)
                else:
                    run = new_p.add_run(part)
                    set_font(run, "KoPub돋움체_Pro Light", 11)
        current_p = new_p
    return current_p

# 템플릿 내의 태그({{태그}})를 찾아 내용을 삽입하는 함수
def replace_placeholder(doc, placeholder, content, is_inline=False):
    for p in doc.paragraphs:
        if placeholder in p.text:
            if is_inline:
                p.text = p.text.replace(placeholder, content)
            else:
                p.text = p.text.replace(placeholder, "") 
                if content:
                    add_styled_content_at(p, content)    
            return True
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if placeholder in p.text:
                        if is_inline:
                            p.text = p.text.replace(placeholder, content)
                        else:
                            p.text = p.text.replace(placeholder, "")
                            if content:
                                add_styled_content_at(p, content)
                        return True
    return False

def extract_tag(text, tag_name):
    pattern = f"<{tag_name}>(.*?)(?:</{tag_name}>|$)"
    match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
    return match.group(1).strip() if match else ""

# 4. 스마트 라우터
def generate_one_shot(prompt):
    fallback_models = ["gemini-2.5-flash", "gemini-2.5-flash-lite"]
    last_error = ""
    for model_name in fallback_models:
        for attempt in range(2):
            try:
                response = client.models.generate_content(
                    model=model_name, 
                    contents=prompt,
                    config={"temperature": 0.2} 
                )
                if response and response.text:
                    return response.text.strip(), model_name
            except Exception as e:
                error_msg = str(e)
                last_error = error_msg
                if "429" in error_msg or "RESOURCE_EXHAUSTED" in error_msg:
                    time.sleep(5) 
                    continue
                else:
                    time.sleep(2)
                    continue
    return None, f"생성 실패 (에러: {last_error[:50]})"

# ★ 외부에서 호출하는 실행 함수
def run_virtual_firm(spec_file, doc_template, target_corp, ir_data, business_status):
    if not spec_file:
        st.error("분석을 위한 기술 명세서(PDF)가 필요합니다.")
        return

    st.subheader(f"🏢 {target_corp if target_corp else 'Virtual Firm'} 심층 창업/사업화 보고서 생성 중...")
    
    with st.spinner("🚀 완벽한 비즈니스 이전을 위한 심층 사업 타당성 및 '기술가치평가 산출'을 진행 중입니다... (약 20~30초 소요)"):
        try:
            tech_text = extract_text_from_pdf(spec_file)
            ir_text = extract_text_from_pdf(ir_data) if ir_data else ""
            context = f"대상기업: {target_corp}\n기업IR자료: {ir_text}\n사업현황: {business_status}" if (target_corp or ir_text or business_status) else "특정 기업 정보 없음"
            
            # 🔥 [Section 3 부분 기술가치평가 산출 수식 및 근거 강제 지시 추가]
            prompt = f"""당신은 부산대학교 기술지주회사의 최고급 비즈니스 아키텍트(Business Architect)이자 '공인 기술가치평가사'입니다. 
            우리의 핵심 철학은 "단순히 '기술(Technology)'을 이전하는 것이 아니라, 완벽하게 기획된 '하나의 사업(Business)' 자체를 이전하거나 창업한다"는 'Virtual Firm' 개념입니다.
            제공된 특허 데이터를 분석하여, 대상 기업인 '{target_corp}'(또는 잠재적 투자자/연구자)에게 이 기술이 얼마나 폭발적인 수익을 낼 수 있는 '사업 아이템'인지 증명하는 심층 보고서를 작성하세요.

            [작성 규칙 - 절대 엄수]
            1. 마크다운 표(| 구분 | 내용 | 등)를 절대 사용하지 마세요. (시스템 에러 발생)
            2. 비교 분석, 재무 수치, 로드맵 등은 반드시 **소제목(##)과 개조식(Bullet points, -, *)을 활용한 깔끔한 텍스트**로만 서술하세요.
            3. 모든 전략과 재무적 예측에는 반드시 논리적이고 객관적인 '산출 근거'를 명확히 제시해야 합니다.
            4. 응답은 반드시 아래 5개의 <태그>를 순서대로 모두 열고 닫아야 합니다.

            <tech_title>이 기술을 기반으로 한 가상 기업(사업)의 핵심을 나타내는 매력적인 비즈니스 명칭 (20자 내외 한 줄)</tech_title>
            
            <section_1>
            (이곳에 'Ⅰ. 기술 개요'를 1,000자 이상 심층 작성하세요. 
            - 기술의 작동 원리와 이 기술이 가지는 '기술적 우위성' 및 '경쟁력'
            - 이 기술이 어떠한 비즈니스적 해자(Moat)를 구축할 수 있는지 분석)
            </section_1>
            
            <section_2>
            (이곳에 'Ⅱ. 문제점 및 해결 방안'을 800자 이상 심층 작성하세요. 
            - 타겟 시장(또는 산업)이 겪고 있는 치명적인 한계점(Pain points)
            - 본 기술이 이를 어떻게 혁신적으로 해결하는지 대조 및 산업적 파급 효과 분석)
            </section_2>
            
            <section_3>
            (이곳에 'Ⅲ. Scale-up 전략 및 심층 재무 계획'을 1,500자 이상 아주 상세히 작성하세요. 
            반드시 아래 3가지 파트를 소제목(##)으로 구분하여 작성하세요.
            ## 1. 사업화 로드맵
            - 상용화 및 시장 진입을 위한 마일스톤별 타임라인
            ## 2. 시장 규모 추정 및 예상 매출액
            - TAM-SAM-SOM 추정치 및 논리적 산출 근거
            - 사업화 이후 3~5년간 발생할 예상 매출액/수익 예측치와 근거 (예상 단가, 점유율 등 명시)
            ## 3. 기술가치평가 (Technology Valuation)
            - 본 기술의 최종적인 '경제적 가치 금액(Valuation)' 명시
            - 평가에 적용된 방법론 명시 (예: 수익접근법, 로열티공제법 등)
            - **[가액 산출 과정 상세 설명]:** AI가 추정한 수치를 활용하여 (추정 매출액의 현가 합계 × 로열티율 × 기술기여도 등) 실제 기술가치평가 산출 수식을 풀어서 단계별로 작성하고 계산 과정을 보여주세요. 각 변수(로열티율, 기여도 등)를 그렇게 설정한 시장/재무적 근거도 반드시 포함하세요.
            )
            </section_3>
            
            <section_4>
            (이곳에 'Ⅳ. Virtual Firm 활용 (최종 사업화 제안)'을 1,500자 이상 아주 상세히 작성하세요. 
            반드시 아래 4가지 파트를 순서대로 빠짐없이 소제목(##)과 개조식으로 포함하여 작성해야 합니다.
            
            ## 1. 3C 분석
            - Company (자사/본 기술의 역량)
            - Competitor (경쟁사 및 대체재 분석)
            - Customer (타겟 고객 및 시장 니즈)
            
            ## 2. SWOT 분석
            - Strength (강점)
            - Weakness (약점)
            - Opportunity (기회)
            - Threat (위협)
            
            ## 3. Lean Canvas (핵심 요약)
            - 문제(Problem), 해결책(Solution), 고유가치제안(UVP), 압도적 경쟁우위(Unfair Advantage), 고객군(Customer Segments), 핵심 지표(Key Metrics), 채널(Channels), 비용구조(Cost Structure), 수익원(Revenue Streams)
            
            ## 4. 최종 제안: 비즈니스 이전 vs 직접 창업
            - 앞선 3C, SWOT, Lean Canvas 분석과 Ⅲ장의 기술가치평가 금액을 종합하여 결론을 내립니다.
            - 타겟 기업에게 "완성된 신사업 모델로서의 '기술이전(M&A급)'"을 제안할지, 아니면 해당 연구자에게 직접 '기술 기반 창업(Startup)'을 제안할지 명확히 선택하고, 정량/정성적 근거를 바탕으로 강력하게 설득하세요.
            )
            </section_4>
            
            [분석 대상 데이터]
            {context}
            명세서 내용: {tech_text}"""

            raw_response, used_model = generate_one_shot(prompt)
            
            if not raw_response:
                st.error(used_model) 
                return

            st.toast(f"✅ 심층 비즈니스 및 기술가치평가 분석 완료! (사용 모델: {used_model})")

            ai_data = {
                "tech_title": extract_tag(raw_response, "tech_title"),
                "section_1": extract_tag(raw_response, "section_1"),
                "section_2": extract_tag(raw_response, "section_2"),
                "section_3": extract_tag(raw_response, "section_3"),
                "section_4": extract_tag(raw_response, "section_4"),
            }

            clean_title = ai_data.get("tech_title", "본 기술").replace('#', '').replace('*', '').strip()
            if not clean_title: clean_title = "심층 사업화 보고서"

            doc = None
            used_template = False

            if doc_template and doc_template.name.lower().endswith('.docx'):
                doc_template.seek(0)
                doc = Document(doc_template)
                used_template = True
            elif os.path.exists(DEFAULT_WORD_TEMPLATE):
                doc = Document(DEFAULT_WORD_TEMPLATE)
                used_template = True
                st.info("ℹ️ 업로드된 양식이 없어 시스템에 내장된 **기본 템플릿**을 자동 적용했습니다.")

            if used_template:
                replace_placeholder(doc, "{{tech_title}}", clean_title, is_inline=True)
                replace_placeholder(doc, "{{section_1}}", ai_data.get("section_1", ""))
                replace_placeholder(doc, "{{section_2}}", ai_data.get("section_2", ""))
                replace_placeholder(doc, "{{section_3}}", ai_data.get("section_3", ""))
                replace_placeholder(doc, "{{section_4}}", ai_data.get("section_4", ""))
                
                st.success("✨ '기술가치평가 산출 과정'이 완벽하게 포함된 Virtual Firm 보고서 생성이 완료되었습니다!")
            else:
                st.warning("⚠️ 지정된 기본 템플릿 파일을 찾을 수 없어 백지 양식으로 생성합니다. (폴더 내에 'default_vf_template.docx'를 넣어주세요)")
                doc = Document()
                title_p = doc.add_paragraph("\n\n")
                title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_run = title_p.add_run(f"Virtual Firm 심층 사업화 전략 보고서\n\n[{clean_title}]")
                set_font(title_run, "KoPub돋움체_Pro Bold", 16)
                
                sections = [("Ⅰ. 기술 개요", "section_1"), ("Ⅱ. 문제점 및 해결 방안", "section_2"), 
                            ("Ⅲ. Scale-up 및 기술가치평가", "section_3"), ("Ⅳ. Virtual Firm 활용 (사업화 최종 제안)", "section_4")]

                for title_text, key in sections:
                    h_p = doc.add_paragraph()
                    set_font(h_p.add_run(f"\n{title_text}"), "KoPub돋움체_Pro Bold", 14)
                    
                    content = ai_data.get(key, "")
                    if content:
                        add_styled_content_at(h_p, content)
                    else:
                        p = doc.add_paragraph()
                        set_font(p.add_run("생성된 내용이 없습니다."), "KoPub돋움체_Pro Light", 11)

            doc_io = io.BytesIO()
            doc.save(doc_io)
            
            st.download_button(
                label="📥 심층 Virtual Firm 보고서 다운로드 (클릭)", 
                data=doc_io.getvalue(), 
                file_name="Virtual_Firm_Master_Report.docx", 
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
        except Exception as e:
            st.error(f"Virtual Firm 생성 중 오류 발생: {e}")
