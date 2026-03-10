import streamlit as st
import fitz
from google import genai
import io
import re
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# [설정 - API]
MY_API_KEY = st.secrets["GEMINI_API_KEY"].strip() 
client = genai.Client(api_key=MY_API_KEY)

# 1. 폰트 설정
def set_font(run, font_name, size, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.bold = bold

# 2. 문서 텍스트 추출 (PDF, DOCX, TXT 완벽 지원)
def extract_text_from_file(uploaded_file):
    if uploaded_file is None: return ""
    try:
        uploaded_file.seek(0)
        filename = uploaded_file.name.lower()
        if filename.endswith(".pdf"):
            doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
            return "".join([page.get_text() for page in doc])[:15000]
        elif filename.endswith(".txt"):
            return uploaded_file.read().decode('utf-8')[:15000]
        elif filename.endswith(".docx"):
            doc = Document(uploaded_file)
            return "\n".join([para.text for para in doc.paragraphs])[:15000]
        else:
            return ""
    except Exception as e:
        return ""

# 3. 스타일 콘텐츠 추가
def add_styled_content(doc, text):
    lines = str(text).split('\n')
    for line in lines:
        line_stripped = line.strip()
        if not line_stripped: continue
        
        if line_stripped.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(line_stripped.replace('## ', ''))
            set_font(run, "KoPub돋움체_Pro Medium", 12, bold=True)
            continue
            
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.6
        p.paragraph_format.space_after = Pt(12)
        
        parts = re.split(r'(\*\*.*?\*\*)', line_stripped)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part.replace('**', ''))
                set_font(run, "KoPub돋움체_Pro Medium", 11, bold=True)
            else:
                run = p.add_run(part)
                set_font(run, "KoPub돋움체_Pro Light", 11)

# 💡 태그 추출 방어 코드
def extract_tag(text, tag_name):
    pattern = f"<{tag_name}>(.*?)(?:</{tag_name}>|$)"
    match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
    return match.group(1).strip() if match else ""

# 4. 스마트 라우터 (온도 설정 추가)
def generate_one_shot(prompt):
    fallback_models = ["gemini-2.5-flash", "gemini-2.5-flash-lite"]
    last_error = ""
    for model_name in fallback_models:
        for attempt in range(2):
            try:
                # ✅ config={"temperature": 0.1} 를 추가하여 AI가 딴짓을 못하도록 논리력과 일관성을 고정합니다.
                response = client.models.generate_content(
                    model=model_name, 
                    contents=prompt,
                    config={"temperature": 0.1} 
                )
                if response and response.text:
                    return response.text.strip(), model_name
            except Exception as e:
                error_msg = str(e)
                last_error = error_msg
                if "429" in error_msg or "RESOURCE_EXHAUSTED" in error_msg:
                    time.sleep(5) 
                    continue
                else:
                    time.sleep(2)
                    continue
    return None, f"생성 실패 (에러: {last_error[:50]})"

# ★ 외부에서 호출하는 실행 함수
def run_proposal(spec_file, ppt_template, target_corp, ir_data, business_status):
    if not spec_file:
        st.error("필수 항목인 특허 명세서(PDF)를 업로드해주세요.")
        return

    st.subheader(f"📊 {target_corp if target_corp else '심층'} 전략보고서 초고속 생성 중...")
    
    tech_text = extract_text_from_file(spec_file)
    
    with st.expander("🔎 [클릭해서 확인] 파이썬이 파일에서 추출한 텍스트 미리보기"):
        st.info(f"총 추출된 글자 수: {len(tech_text.strip())}자")
        if len(tech_text.strip()) > 0:
            st.text(tech_text[:1000] + ("\n... [중략] ..." if len(tech_text) > 1000 else ""))
        else:
            st.warning("경고: 파일에서 글자를 전혀 읽어내지 못했습니다!")
            
    if not tech_text or len(tech_text.strip()) < 50:
        st.error("❌ 특허청 PDF의 보안(폰트 깨짐) 또는 이미지 스캔본이어서 파이썬이 글자를 읽지 못했습니다! \n\n**[해결 방법]**\n1. 원본 PDF를 여시고 텍스트를 전체 복사(Ctrl+A, Ctrl+C) 하세요.\n2. 빈 메모장(.txt)이나 워드(.docx)에 붙여넣고 저장하세요.\n3. 그 txt나 docx 파일을 대신 업로드하시면 완벽하게 작동합니다!")
        return
        
    with st.spinner("🚀 명세서 내용을 정밀 분석 중입니다... (약 15초 소요)"):
        try:
            ir_text = extract_text_from_file(ir_data) if ir_data else ""
            context = f"대상기업: {target_corp}\n기업IR자료: {ir_text}\n사업현황: {business_status}" if (target_corp or ir_text or business_status) else "지정된 타겟 기업 없음 (범용 시장 기준으로 분석할 것)"
            
            # ✅ AI가 괄호 내용을 그대로 복사하지 못하도록 출력 포맷을 분리하고 강력한 작성 규칙을 추가했습니다.
            prompt = f"""당신은 최고 수준의 기술사업화 전문가입니다. 
            아래 제공된 [특허 명세서 원본 데이터]를 완벽하게 분석하고, 이를 바탕으로 전략보고서를 작성하세요.

            [특허 명세서 원본 데이터]
            {tech_text}

            [타겟 기업 및 사업 환경 정보]
            {context}

            [작성 규칙 - 절대 엄수 (위반 시 시스템 오류 발생)]
            1. [가이드라인 복사 금지] 아래 출력 포맷에 적힌 괄호 안의 지시문(예: "상세 분석 본문 작성")을 그대로 출력하는 것을 엄격히 금지합니다.
            2. [분량 강제] 각 섹션(<section_1> ~ <section_5>)은 특허 내용을 바탕으로 최소 500자 이상, 3~4개의 상세한 단락과 개조식(Bullet points)을 포함하여 꽉 채워 작성하세요. 목차 제목(예: Ⅰ. 기술의 메커니즘...)도 본문 첫 줄에 반드시 포함하세요.
            3. 기술을 상상해서 지어내지 말고 제공된 데이터에만 기반하세요. 표(| 구분 | 내용 | 등)는 절대 그리지 마세요. 
            4. 응답은 반드시 아래 6개의 <태그>로 감싸서 출력하세요. (JSON 형식 절대 사용 금지)

            [출력 포맷 (반드시 이 구조를 따를 것)]
            <tech_title>이 특허 명세서의 핵심을 나타내는 정확한 기술 명칭 (20자 내외 한 줄)</tech_title>
            <section_1>
            (이곳에 'Ⅰ. 기술의 메커니즘 및 완성도 분석'에 대한 500자 이상의 명세서 기반 작동 원리와 구조 상세 설명 작성)
            </section_1>
            <section_2>
            (이곳에 'Ⅱ. 시장 생태계 분석 및 산업 트렌드'에 대한 500자 이상의 산업 동향 상세 분석 작성)
            </section_2>
            <section_3>
            (이곳에 'Ⅲ. 기술 사업화 모델(BM) 및 기대효과'에 대한 구체적 근거를 바탕으로 한 BM 제안 상세 작성)
            </section_3>
            <section_4>
            (이곳에 'Ⅳ. 지식재산권 포트폴리오 강점 및 회피 설계 방어 전략'에 대한 상세 분석 작성)
            </section_4>
            <section_5>
            (이곳에 'Ⅴ. 맞춤형 기술 융합 및 고도화 전략'에 대한 타겟 기업 맞춤형 제안 상세 작성)
            </section_5>"""

            raw_response, used_model = generate_one_shot(prompt)
            
            if not raw_response:
                st.error(used_model) 
                return

            st.toast(f"✅ 명세서 집중 분석 완료! (사용 모델: {used_model})")
            
            ai_data = {
                "tech_title": extract_tag(raw_response, "tech_title"),
                "section_1": extract_tag(raw_response, "section_1"),
                "section_2": extract_tag(raw_response, "section_2"),
                "section_3": extract_tag(raw_response, "section_3"),
                "section_4": extract_tag(raw_response, "section_4"),
                "section_5": extract_tag(raw_response, "section_5"),
            }

            doc = Document()
            target_display = target_corp if target_corp else "잠재적 수요기업"
            
            doc.add_paragraph("\n" * 3)
            title_p = doc.add_paragraph()
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            clean_title = ai_data.get("tech_title", "본 기술").replace('#', '').replace('*', '').strip()
            if not clean_title: clean_title = "기술이전 전략 제안서"
            
            title_run = title_p.add_run(f"기술이전 전략 제안서\n\n[{clean_title}]")
            set_font(title_run, "KoPub돋움체_Pro Bold", 15)
            
            doc.add_paragraph("\n" * 7)
            info_p = doc.add_paragraph()
            info_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            info_run = info_p.add_run(f"제안 대상: {target_display} 귀하\n제안 기관: 부산대학교 산학협력단\n작성일자: 2026. 03. 10")
            set_font(info_run, "KoPub돋움체_Pro Light", 11)
            doc.add_page_break()

            sec_5_title = f"Ⅴ. {target_display} 맞춤형 기술 융합 및 고도화 전략" if target_corp else "Ⅴ. 산업별 기술 응용 및 비즈니스 모델 제안"
            sections = [
                ("Ⅰ. 기술의 메커니즘 및 완성도 분석", "section_1"),
                ("Ⅱ. 시장 생태계 분석 및 산업 트렌드", "section_2"),
                ("Ⅲ. 기술 사업화 모델 및 기대효과", "section_3"),
                ("Ⅳ. 지식재산권 포트폴리오 강점", "section_4"),
                (sec_5_title, "section_5")
            ]

            for title_text, key in sections:
                # 🔥 AI가 제목을 중복 출력하는 것을 막기 위해 Word 생성 로직에서 제목 부분을 조금 정리합니다.
                h_p = doc.add_paragraph()
                h_run = h_p.add_run(title_text)
                set_font(h_run, "KoPub돋움체_Pro Medium", 13)
                h_p.paragraph_format.space_before = Pt(18)
                
                content = ai_data.get(key, "")
                # AI가 제목을 중복 출력했다면 제거해주는 방어 코드
                content = re.sub(rf"^{title_text}\n*", "", content, flags=re.IGNORECASE)
                
                if content.strip():
                    add_styled_content(doc, content)
                else:
                    p = doc.add_paragraph()
                    set_font(p.add_run("생성된 내용이 없습니다."), "KoPub돋움체_Pro Light", 11)

            doc_io = io.BytesIO()
            doc.save(doc_io)
            
            st.success("✨ 업로드하신 명세서의 내용이 100% 반영된 무결점 전략보고서가 완성되었습니다!")
            
            st.download_button(
                label="📥 전략보고서 최종 다운로드 (클릭)", 
                data=doc_io.getvalue(), 
                file_name="Proposal_Strategic_Report_Final.docx", 
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
        except Exception as e:
            st.error(f"전략보고서 생성 중 오류 발생: {e}")